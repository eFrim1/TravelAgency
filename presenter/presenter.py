from model.repository.package_repository import PackageRepository
from model.repository.client_repository import ClientRepository
from model.repository.bought_package_repository import BoughtPackageRepository 
from .interface import ITravelAgencyGUI
from docx import Document
from model.models import Package, Client, BoughtPackage

class Presenter:
    def __init__(self, gui: ITravelAgencyGUI):
        if not isinstance(gui, ITravelAgencyGUI):
            raise TypeError("Expected an instance of ITravelAgencyGUI")
        self.gui = gui
        self.package_repository = PackageRepository()
        self.client_repository = ClientRepository()
        self.bought_package_repository = BoughtPackageRepository()
    
    def UI_initialized(self):

        self.gui.display_packages(*self.package_model_list_to_parameters_lists(self.package_repository.get_all()))
        self.gui.display_clients(*self.client_model_list_to_parameters_lists_clients(self.client_repository.get_all()))
        self.gui.display_bought_packages(*self.bought_package_model_list_to_parameters_lists_bought_packages(self.bought_package_repository.get_all()))
    
    def row_selected(self, table_type, row):
        if row == -1:
            return
        if table_type == "package":
            package_data = self.gui.get_row_selected_data(table_type)
            package = self.package_repository.get_by_id(package_data['id'])
            self.gui.display_package_form_data(package)
            self.gui.display_images(package.image_1, package.image_2, package.image_3)
        elif table_type == "client":
            client_data = self.gui.get_row_selected_data(table_type)
            client = self.client_repository.get_by_id(client_data['id'])
            self.gui.display_client_form_data(client)
        elif table_type == "bought_package":
            bought_package_data = self.gui.get_row_selected_data(table_type)
            bought_package = self.bought_package_repository.get_by_id(bought_package_data['id'])
            self.gui.display_bought_package_form_data(bought_package)

    def package_model_list_to_parameters_lists(self, packages):
        packages_id = []
        packages_dest = []
        packages_period = []
        packages_price = []
        packages_image1 = []
        packages_image2 = []
        packages_image3 = []
        for package in packages:
            packages_id.append(package.id)
            packages_dest.append(package.destination)
            packages_period.append(package.period)
            packages_price.append(package.price)
            packages_image1.append(package.image_1)
            packages_image2.append(package.image_2)
            packages_image3.append(package.image_3)
        return packages_id, packages_dest, packages_period, packages_price, packages_image1, packages_image2, packages_image3

    def client_model_list_to_parameters_lists_clients(self, clients):
        clients_id = []
        clients_name = []
        clients_email = []
        clients_phone = []
        for client in clients:
            clients_id.append(client.id)
            clients_name.append(client.name)
            clients_email.append(client.email)
            clients_phone.append(client.phone)
        return clients_id, clients_name, clients_email, clients_phone

    def bought_package_model_list_to_parameters_lists_bought_packages(self, bought_packages):
        bought_packages_id = []
        bought_packages_package_id = []
        bought_packages_client_id = []
        bought_packages_bought_date = []
        bought_packages_departure = []
        bought_packages_arrival = []
        for bought_package in bought_packages:
            bought_packages_id.append(bought_package.id)
            bought_packages_package_id.append(bought_package.package_id)
            bought_packages_client_id.append(bought_package.client_id)
            bought_packages_bought_date.append(bought_package.bought_date)
            bought_packages_departure.append(bought_package.departure)
            bought_packages_arrival.append(bought_package.arrival)
        return bought_packages_id, bought_packages_package_id, bought_packages_client_id, bought_packages_bought_date, bought_packages_departure, bought_packages_arrival


    def add_item(self):
        table_type = self.gui.get_table_type()
        if table_type == "package":
            package_data = self.gui.get_package_form_data()
            package = Package(**package_data)
            self.package_repository.add(package)
            packages = self.package_repository.get_all()
            self.gui.display_packages(*self.package_model_list_to_parameters_lists(packages))
        elif table_type == "client":
            client_data = self.gui.get_client_form_data()
            client = Client(**client_data)
            self.client_repository.add(client)
            clients = self.client_repository.get_all()
            self.gui.display_clients(*self.client_model_list_to_parameters_lists_clients(clients))
        elif table_type == "bought_package":
            bought_package_data = self.gui.get_bought_package_form_data()
            bought_package = BoughtPackage(**bought_package_data)
            self.bought_package_repository.add(bought_package)
            bought_packages = self.bought_package_repository.get_all()
            self.gui.display_bought_packages(*self.bought_package_model_list_to_parameters_lists_bought_packages(bought_packages))

    
    def update_item(self):
        table_type = self.gui.get_table_type()
        if table_type == "package":
            package_data = self.gui.get_package_form_data()
            package = Package(**package_data)
            self.package_repository.update(package)
            packages = self.package_repository.get_all()
            self.gui.display_packages(*self.package_model_list_to_parameters_lists(packages))
        elif table_type == "client":
            client_data = self.gui.get_client_form_data()
            client = Client(**client_data)
            self.client_repository.update(client)
            clients = self.client_repository.get_all()
            self.gui.display_clients(*self.client_model_list_to_parameters_lists_clients(clients))
        elif table_type == "bought_package":
            bought_package_data = self.gui.get_bought_package_form_data()
            bought_package = BoughtPackage(**bought_package_data)
            self.bought_package_repository.update(bought_package)
            bought_packages = self.bought_package_repository.get_all()
            self.gui.display_bought_packages(*self.bought_package_model_list_to_parameters_lists_bought_packages(bought_packages))

    
    def delete_item(self):
        table_type = self.gui.get_table_type()
        if table_type == "package":
            package_data = self.gui.get_package_form_data()
            package = Package(**package_data)
            self.package_repository.delete(package.id)
            packages = self.package_repository.get_all()
            self.gui.display_packages(*self.package_model_list_to_parameters_lists(packages))
        elif table_type == "client":
            client_data = self.gui.get_client_form_data()
            client = Client(**client_data)
            self.client_repository.delete(client.id)
            clients = self.client_repository.get_all()
            self.gui.display_clients(*self.client_model_list_to_parameters_lists_clients(clients))
        elif table_type == "bought_package":
            bought_package_data = self.gui.get_bought_package_form_data()
            bought_package = BoughtPackage(**bought_package_data)
            self.bought_package_repository.delete(bought_package.id)
            bought_packages = self.bought_package_repository.get_all()
            self.gui.display_bought_packages(*self.bought_package_model_list_to_parameters_lists_bought_packages(bought_packages))


    def export_to_csv(self):
        table_type = self.gui.get_table_type()
        if table_type == "package":
            packages = self.package_repository.get_all()
            with open("packages.csv", "w") as file:
                file.write("ID,Destination,Period,Price,Image_1,Image_2,Image_3\n")
                for package in packages:
                    file.write(f"{package.id},{package.destination},{package.period},{package.price},{package.image_1},{package.image_2},{package.image_3}\n")
        elif table_type == "client":
            clients = self.client_repository.get_all()
            with open("clients.csv", "w") as file:
                file.write("ID,Name,Email,Phone\n")
                for client in clients:
                    file.write(f"{client.id},{client.name},{client.email},{client.phone}\n")
        elif table_type == "bought_package":
            bought_packages = self.bought_package_repository.get_all()
            with open("bought_packages.csv", "w") as file:
                file.write("ID,Package ID,Client ID,Bought Date,Departure,Arrival\n")
                for bought_package in bought_packages:
                    file.write(f"{bought_package.id},{bought_package.package_id},{bought_package.client_id},{bought_package.bought_date},{bought_package.departure},{bought_package.arrival}\n")
        
    def export_to_doc(self):
        table_type = self.gui.get_table_type()
        document = Document()
        if table_type == "package":
            packages = self.package_repository.get_all()
            table = document.add_table(rows=1, cols=7)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Destination'
            hdr_cells[2].text = 'Period'
            hdr_cells[3].text = 'Price'
            hdr_cells[4].text = 'Image 1'
            hdr_cells[5].text = 'Image 2'
            hdr_cells[6].text = 'Image 3'
            for package in packages:
                row_cells = table.add_row().cells
                row_cells[0].text = str(package.id)
                row_cells[1].text = package.destination
                row_cells[2].text = package.period
                row_cells[3].text = str(package.price)
                row_cells[4].text = package.image_1
                row_cells[5].text = package.image_2
                row_cells[6].text = package.image_3
                document.save('packages.docx')
        elif table_type == "client":
            clients = self.client_repository.get_all()
            table = document.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'Email'
            hdr_cells[3].text = 'Phone'
            for client in clients:
                row_cells = table.add_row().cells
                row_cells[0].text = str(client.id)
                row_cells[1].text = client.name
                row_cells[2].text = client.email
                row_cells[3].text = client.phone
                document.save('clients.docx')
        elif table_type == "bought_package":
            bought_packages = self.bought_package_repository.get_all()
            table = document.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Package ID'
            hdr_cells[2].text = 'Client ID'
            hdr_cells[3].text = 'Bought Date'
            hdr_cells[4].text = 'Departure'
            hdr_cells[5].text = 'Arrival'
            for bought_package in bought_packages:
                row_cells = table.add_row().cells
                row_cells[0].text = str(bought_package.id)
                row_cells[1].text = str(bought_package.package_id)
                row_cells[2].text = str(bought_package.client_id)
                row_cells[3].text = bought_package.bought_date
                row_cells[4].text = bought_package.departure
                row_cells[5].text = bought_package.arrival
                document.save('bought_packages.docx')

    def search_bought_packages(self):
        search_input = self.gui.get_search_client_input()
        bought_packages = self.bought_package_repository.get_by_client_name(search_input)
        self.gui.display_bought_packages(*self.bought_package_model_list_to_parameters_lists_bought_packages(bought_packages))
    
    def filter_packages(self):
        (filter_type, filter_input) = self.gui.get_filter_input_combo()
        filter_input = filter_input.strip()
        filter_type = filter_type.strip().lower()
        compare = "equal"
        if filter_input:
            if filter_input[0] == "<":
                filter_input = filter_input[1:]
                compare = "less"
            elif filter_input[0] == ">":
                filter_input = filter_input[1:]
                compare = "greater"
        packages = self.package_repository.filter(filter_str=filter_input, filter_type=filter_type, compare_type=compare)
        self.gui.display_packages(*self.package_model_list_to_parameters_lists(packages))

    
    def sort_packages(self):
        sort_combo = self.gui.get_sort_combo().strip().lower()
        packages = self.package_repository.sort(sort_combo)
        self.gui.display_packages(*self.package_model_list_to_parameters_lists(packages))


